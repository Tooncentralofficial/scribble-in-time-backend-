from django.db import models
from django.conf import settings
from django.utils import timezone

# Import the custom User model from custom_user.py
from .custom_user import User

class Conversation(models.Model):
    STATUS_CHOICES = [
        ('pending_ai', 'Pending AI Response'),
        ('awaiting_admin', 'Awaiting Admin Response'),
        ('admin_responded', 'Admin Has Responded'),
        ('resolved', 'Resolved'),
    ]
    
    user = models.ForeignKey('scribble.User', on_delete=models.SET_NULL, null=True, blank=True, related_name='conversations')
    session_key = models.CharField(max_length=40, null=True, blank=True, help_text='Session key for unauthenticated users')
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    is_active = models.BooleanField(default=True)
    ai_enabled = models.BooleanField(default=True)
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='pending_ai')

    class Meta:
        ordering = ['-updated_at']

class Message(models.Model):
    SENDER_CHOICES = [
        ('user', 'User'),
        ('ai', 'AI'),
        ('admin', 'Admin')
    ]
    
    conversation = models.ForeignKey(Conversation, on_delete=models.CASCADE, related_name='messages')
    content = models.TextField()
    sender = models.CharField(max_length=10, choices=SENDER_CHOICES)
    created_at = models.DateTimeField(auto_now_add=True)
    is_read = models.BooleanField(default=False)

    class Meta:
        ordering = ['created_at']

class Document(models.Model):
    DOCUMENT_TYPES = [
        ('knowledge', 'Knowledge Base'),
        ('faq', 'FAQ'),
        ('other', 'Other')
    ]
    
    title = models.CharField(max_length=255)
    document_type = models.CharField(max_length=20, choices=DOCUMENT_TYPES, default='knowledge')
    file = models.FileField(upload_to='documents/')
    uploaded_by = models.ForeignKey('scribble.User', on_delete=models.SET_NULL, null=True, related_name='uploaded_documents')
    uploaded_at = models.DateTimeField(auto_now_add=True)
    is_active = models.BooleanField(default=True)

    def __str__(self):
        return self.title

class AdminSettings(models.Model):
    user = models.OneToOneField('scribble.User', on_delete=models.CASCADE, related_name='admin_settings')
    ai_enabled = models.BooleanField(default=True)
    auto_response = models.BooleanField(default=True)
    response_timeout = models.PositiveIntegerField(default=30)  # in minutes
    updated_at = models.DateTimeField(auto_now=True)

    def __str__(self):
        return f"Settings for {self.user.email}"

import os

class KnowledgeDocument(models.Model):
    file = models.FileField(upload_to='knowledge_base/')
    title = models.CharField(max_length=255, blank=True)
    content = models.TextField(blank=True, help_text='Processed content for vector search')
    uploaded_at = models.DateTimeField(auto_now_add=True)
    is_processed = models.BooleanField(default=False)
    processing_error = models.TextField(blank=True, null=True, help_text='Error message if processing failed')
    file_size = models.PositiveIntegerField(null=True, blank=True)
    file_type = models.CharField(max_length=50, blank=True)

    def save(self, *args, **kwargs):
        if not self.title and self.file:
            self.title = os.path.basename(self.file.name)
        if self.file:
            self.file_size = self.file.size
            self.file_type = os.path.splitext(self.file.name)[1].lower()
        super().save(*args, **kwargs)

    def process_document(self):
        """Process the document content for vector search"""
        try:
            import PyPDF2
            import docx
            import textract
            
            if self.file_type == '.pdf':
                pdf_reader = PyPDF2.PdfReader(self.file)
                self.content = "\n".join([page.extract_text() for page in pdf_reader.pages])
            elif self.file_type == '.docx':
                doc = docx.Document(self.file)
                self.content = "\n".join([para.text for para in doc.paragraphs])
            else:
                # Fallback to textract for other formats
                self.content = textract.process(self.file.path).decode('utf-8')
            
            self.is_processed = True
            self.save()
            return True
        except Exception as e:
            print(f"Error processing document {self.id}: {str(e)}")
            return False

    def __str__(self):
        return self.title or f"Document {self.id}"

class MemoirFormSubmission(models.Model):
    """Model for memoir form submissions"""
    GENDER_CHOICES = [
        ('male', 'Male'),
        ('female', 'Female'),
        ('other', 'Other'),
        ('prefer_not_to_say', 'Prefer not to say'),
    ]
    
    AUDIENCE_CHOICES = [
        ('family_friends', 'Family and Friends'),
        ('public', 'Public'),
        ('specific_group', 'Specific Group'),
    ]
    
    # Personal Information
    first_name = models.CharField(max_length=100)
    last_name = models.CharField(max_length=100)
    email = models.EmailField()
    phone_number = models.CharField(max_length=20)
    gender = models.CharField(max_length=20, choices=GENDER_CHOICES)
    
    # Memoir Details
    theme = models.CharField(max_length=200, help_text="Overall theme of the memoir")
    subject = models.CharField(max_length=200, help_text="Subject of the memoir")
    main_themes = models.TextField(help_text="Main themes to cover in the memoir")
    key_life_events = models.TextField(help_text="Key life events to include")
    audience = models.CharField(max_length=20, choices=AUDIENCE_CHOICES)
    
    # Metadata
    submitted_at = models.DateTimeField(auto_now_add=True)
    is_processed = models.BooleanField(default=False)
    processing_notes = models.TextField(blank=True, null=True)
    
    class Meta:
        ordering = ['-submitted_at']
        verbose_name = "Memoir Form Submission"
        verbose_name_plural = "Memoir Form Submissions"
    
    def __str__(self):
        return f"Memoir Submission - {self.first_name} {self.last_name} ({self.email})"
    
    @property
    def full_name(self):
        return f"{self.first_name} {self.last_name}"
    
    def get_audience_display_name(self):
        return dict(self.AUDIENCE_CHOICES).get(self.audience, self.audience)
